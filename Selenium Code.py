from selenium import webdriver
from selenium.common.exceptions import NoSuchElementException
driver=webdriver.Chrome()          #For Selenium version above 4.10.0 , we don't need to install chrome webdriver manually


driver.get("https://github.com/Git-Vivek-Gupta/Web-scrapping-project") #URL of the website for web-scrapping


def exist(path):                   #Function for taking data from CSS selector path of the Element as string
    try:
        present=driver.find_element("css selector", path).text
    except NoSuchElementException:
        present=''
    return present
def existx(path):                  #Function for taking data from  xpath of the Element as string
    try:
        present=driver.find_element("xpath", path).text
    except NoSuchElementException:
        present=''
    return present

p= exist("p[dir='auto']")          #Giving the Rel CSS selector path of the Selected Text as argument to the function 'exist'

driver.close()                     #the webpage will be closed automatically


from docx import Document          #docx is the python library to work with docx files
from docx.shared import Inches     #Base class for length constructor classes Inches, Cm, Mm, Px, and Emu. Behaves as an int count of English Metric Units, 914,400 to the inch, 36,000 to the mm.
doc=Document()
doc.add_paragraph(str(p))          #the argument for add_paragraph function must be converted into string

doc.save(str("Vivek'repository")+'.docx')       #give the name of the file to save
